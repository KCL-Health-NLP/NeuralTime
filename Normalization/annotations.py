from docx import Document
import pandas as pd
from docx.enum.text import WD_COLOR_INDEX

def format_file(document_path, output_path, annotations):

   """

    this function transforms a raw text document into a highlighted word document for annotation purposes

   :param document_path: the path to the text document
   :param relative_timexes: a dataframe containing the relative time annotations which will be annotated
    and the absolute time annotations used for reference. attributes should contain : start, end and absolute (boolean)
   :return: nothing, saves a docx file
   """

   text = open(document_path, 'r').read()
   document = Document()
   paragraph = document.add_paragraph()

   text_counter = 0

   for annotation in sorted(annotations.to_dict('records'), key= lambda k : k['start']):
      print(annotation)
      ann_start = annotation['start']
      ann_end = annotation['end']

      # adding the text
      t = text[text_counter:ann_start]
      paragraph.add_run(t)

      # adding the annotation
      ann = text[ann_start:ann_end]
      run = paragraph.add_run(ann)
      # change the style of the annotated text
      if annotation['absolute']:
         run.font.highlight_color = WD_COLOR_INDEX.YELLOW
      else:
         run.font.highlight_color = WD_COLOR_INDEX.BLUE

      # update counter
      text_counter = ann_end

   document.save(output_path)



def create_annotation_table(annotations, output_path):

   """
   creates the table to be filled by the annotator
   :param annotations: all the date and time timexes for the document - with the "absolute" attribute
   :param output_path: path where the table will be saved as an excel file
   :return: the table which was saved
   """

   filtered = annotations[annotations.absolute == False][['ann_text', 'value', 'start', 'end']]
   filtered['IS_RELATIVE'] = [True for i in range(len(filtered))]
   filtered['ANCHOR_TO_ADMISSION'] = [False for i in range(len(filtered))]
   filtered['ANCHOR_TO_DISCHARGE'] = [False for i in range(len(filtered))]
   filtered['ANCHOR_TO_PREVIOUS_TIMEXE'] = [False for i in range(len(filtered))]
   filtered['ANCHOR_TO_PREVIOUS_ABSOLUTE_TIMEXE'] = [False for i in range(len(filtered))]
   filtered['OTHER'] = [None for i in range(len(filtered))]
   filtered['ANCHOR_RELATION'] = [None for i in range(len(filtered))]

   filtered.to_excel(output_path)

   return filtered




timexes = pd.read_excel('date_and_time.xlsx')
doc_path = '../TimeDatasets/i2b2 Data/2012-07-15.original-annotation.release/1.xml.txt'
output = '1.docx'
annotations = timexes[timexes.docname == '1.xml']
print(annotations)
#format_file(doc_path, output, annotations)

create_annotation_table(annotations, '1.xlsx')